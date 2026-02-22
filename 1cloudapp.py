import streamlit as st
import os, io, urllib.parse, time, pandas as pd, re
from datetime import datetime
from google import genai
from docx import Document
from fpdf import FPDF
from supabase import create_client, Client

# ---------------------------------------------------
# 0. UNIVERSAL UI LOCKDOWN
# ---------------------------------------------------
st.set_page_config(
    page_title="VakilDraft",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
    <style>
        #MainMenu {display:none !important;}
        footer {display:none !important;}
        header {visibility:hidden !important;}
        div[data-testid="stToolbar"] {visibility:hidden !important; height:0px !important;}
        .stAppDeployButton {display:none !important;}
        .block-container {padding-top:1rem !important;}
    </style>
""", unsafe_allow_html=True)

# ---------------------------------------------------
# 1. SESSION STATE INIT
# ---------------------------------------------------
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
    st.session_state.user_role = 'user'

if 'final_master' not in st.session_state:
    st.session_state.final_master = ""

if 'draft_history' not in st.session_state:
    st.session_state.draft_history = []

if 'facts_input' not in st.session_state:
    st.session_state.facts_input = ""

if 'selected_model' not in st.session_state:
    st.session_state.selected_model = "Auto-Pilot"

# Research states
if "search_keywords" not in st.session_state:
    st.session_state.search_keywords = []

if "selected_references" not in st.session_state:
    st.session_state.selected_references = []

if "research_court" not in st.session_state:
    st.session_state.research_court = "Kerala High Court"

if "research_period" not in st.session_state:
    st.session_state.research_period = "Last 3 Years"

# ---------------------------------------------------
# 2. LOGIN
# ---------------------------------------------------
if not st.session_state.authenticated:
    st.title("👨‍⚖️ VakilDraft Login")

    with st.form("login_form"):
        u = st.text_input("User")
        p = st.text_input("Password", type="password")

        if st.form_submit_button("Access"):
            creds = st.secrets.get("passwords", {})
            if u in creds and p == creds[u]:
                st.session_state.authenticated = True
                st.session_state.user_role = u.lower()
                st.rerun()
            else:
                st.error("Invalid credentials")

    st.stop()

# ---------------------------------------------------
# 3. STORAGE
# ---------------------------------------------------
SUPABASE_URL = "https://wuhsjcwtoradbzeqsoih.supabase.co"
SUPABASE_KEY = "sb_publishable_02nqexIYCCBaWryubZEkqA_Tw2PqX6m"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

VAULT_PATH = "private_vault"
if not os.path.exists(VAULT_PATH):
    os.makedirs(VAULT_PATH)

# ---------------------------------------------------
# 4. COURT DATA
# ---------------------------------------------------
COURT_DATA = {
    "High Court": ["Writ Petition (Civil)", "Writ Petition (Crl)", "Bail App", "Crl.MC", "Mat.Appeal", "RFA", "RSA"],
    "Family Court": ["OP (Divorce)", "MC (Maintenance)", "GOP (Guardianship)", "OP (Restitution)", "IA (Interim)"],
    "Munsiff Court": ["OS (Original Suit)", "EP (Execution Petition)", "RCP (Rent Control)", "CMA (Misc Appeal)"],
    "DVC (Domestic Violence)": ["DVA (Protection Order)", "Interim Maintenance", "Residence Order"],
    "MC (Magistrate)": ["CMP (Misc Petition)", "ST (Summary Trial)", "CC (Calendar Case)", "Bail Application"],
    "MVOP (Motor Accident)": ["OP (MV) Claim", "Ex-parte Set Aside", "Review Petition"]
}

DIST_SESSIONS_COURT = {
    "Civil": [
        "OS - Original Suit",
        "OP - Original Petition (Divorce / Family)",
        "EA - Execution Application / Petition",
        "MACT - Motor Accident Claim",
        "CMA - Civil Misc. Appeal",
        "Property/Title Dispute",
        "Contract Dispute",
        "Commercial/Company Suit"
    ],
    "Criminal": [
        "Sessions Case - Serious Offences",
        "CRR - Criminal Revision",
        "CRA - Criminal Appeal",
        "CMP - Criminal Misc. Petition",
        "Bail Application",
        "Contempt (Criminal)",
        "NDPS / Special Criminal Cases"
    ]
}

# ---------------------------------------------------
# 5. CORE FUNCTIONS
# ---------------------------------------------------
def perform_replacement(old, new):
    if new and old and "main_editor" in st.session_state:
        updated = st.session_state.main_editor.replace(old, new)
        st.session_state.final_master = updated
        st.session_state.main_editor = updated

def detect_unverified_citation(text):
    patterns = [
        r"\(\d{4}\)\s*\d+\s*SCC",
        r"AIR\s*\d{4}",
        r"\d{4}\s*KHC",
        r"\d{4}\s*Ker\s*LJ"
    ]
    for p in patterns:
        if re.search(p, text):
            return True
    return False

def smart_rotate_draft(prompt, facts, choice):
    projects = st.secrets.get("API_KEYS", [])
    effective_choice = choice if st.session_state.user_role == "admin" else "Auto-Pilot"

    target_model = (
        effective_choice
        if effective_choice != "Auto-Pilot"
        else ("gemini-2.5-pro" if len(facts) > 1200 else "gemini-2.5-flash")
    )

    start_time = time.time()

    for name, key in projects:
        try:
            client = genai.Client(api_key=key)
            res = client.models.generate_content(model=target_model, contents=prompt)
            return res.text, f"{name} ({target_model})", round(time.time() - start_time, 1)
        except:
            continue

    return None, "Offline", 0

def generate_search_keywords(dtype, facts):
    strict_prompt = f"""
Generate exactly 3 short legal issue-based search phrases.
STRICT RULES:
- Do NOT generate case names
- Do NOT generate citations
- Maximum 6 words per phrase
- Output only 3 lines

Petition Type: {dtype}
Facts: {facts}
"""
    result, _, _ = smart_rotate_draft(strict_prompt, facts, st.session_state.selected_model)
    if result:
        lines = [l.strip("- ").strip() for l in result.split("\n") if l.strip()]
        return lines[:3]
    return []

# ---------------------------------------------------
# 6. TOP BAR
# ---------------------------------------------------
col1, col2, col3 = st.columns([6, 2, 1])

with col1:
    st.markdown("## ⚖️ VakilDraft")

with col2:
    st.markdown(f"**Logged in as:** Adv. {st.session_state.user_role.upper()}")

with col3:
    if st.button("🚪 Logout"):
        st.session_state.authenticated = False
        st.session_state.user_role = "user"
        st.rerun()

st.divider()

# ---------------------------------------------------
# 7. COURT SELECTION
# ---------------------------------------------------
c1, c2 = st.columns(2)

with c1:
    court_options = [
        "High Court",
        "Dist & Sessions Court",
        "Family Court",
        "Munsiff Court",
        "DVC (Domestic Violence)",
        "MC (Magistrate)",
        "MVOP (Motor Accident)"
    ]

    court = st.selectbox("Court Level", court_options)

    if court == "Dist & Sessions Court":
        category = st.selectbox("Category", ["Civil", "Criminal"])
        case_types = DIST_SESSIONS_COURT.get(category, [])
        dtype = st.selectbox("Case Type", case_types)
    else:
        dtype = st.selectbox("Petition Type", COURT_DATA.get(court, []))

with c2:
    dists = [
        "Thiruvananthapuram","Kollam","Pathanamthitta","Alappuzha",
        "Kottayam","Idukki","Ernakulam","Thrissur","Palakkad",
        "Malappuram","Kozhikode","Wayanad","Kannur","Kasaragod"
    ]

    if court == "High Court":
        target_dist = "Ernakulam"
        st.text_input("District", value="Ernakulam (High Court of Kerala)", disabled=True)
    else:
        target_dist = st.selectbox("District", dists)

# ---------------------------------------------------
# 8. FACTS INPUT
# ---------------------------------------------------
st.session_state.facts_input = st.text_area(
    "Case Facts:",
    value=st.session_state.facts_input,
    height=150
)

# ---------------------------------------------------
# 8A. LITIGATION RESEARCH MODULE (PROFESSIONAL)
# ---------------------------------------------------
st.divider()
st.subheader("📚 Litigation Research Module (Verified Only)")

st.session_state.research_court = st.radio(
    "Research From:",
    ["Kerala High Court", "Supreme Court", "Both"],
    horizontal=True
)

st.session_state.research_period = st.radio(
    "Research Period:",
    ["Last 3 Years", "Last 5 Years", "No Limit"],
    horizontal=True
)

def generate_google_link(base_site, keywords, period):
    from urllib.parse import quote_plus
    current_year = datetime.now().year
    year_filter = ""

    if period == "Last 3 Years":
        year_filter = f"{current_year}|{current_year-1}|{current_year-2}"
    elif period == "Last 5 Years":
        year_filter = f"{current_year}|{current_year-1}|{current_year-2}|{current_year-3}|{current_year-4}"

    # FIX: Kerala HC uses Indian Kanoon (reliable indexed source)
    if base_site == "highcourt.kerala.gov.in":
        query = f'site:indiankanoon.org "{keywords}" "Kerala High Court" {year_filter}'
    else:
        query = f"site:{base_site} {keywords} judgment {year_filter}"

    return f"https://www.google.com/search?q={quote_plus(query)}"

if st.button("🧠 Generate Official Search Links"):
    if st.session_state.facts_input.strip():
        keywords = generate_search_keywords(dtype, st.session_state.facts_input)
        st.session_state.search_keywords = keywords
    else:
        st.warning("Enter facts first.")

if st.session_state.search_keywords:
    st.markdown("### 🔎 Official Search Links")

    for phrase in st.session_state.search_keywords:

        st.markdown(f"**Search Phrase:** {phrase}")

        if st.session_state.research_court in ["Kerala High Court", "Both"]:
            kerala_link = generate_google_link(
                "highcourt.kerala.gov.in",
                phrase,
                st.session_state.research_period
            )
            st.markdown(f"[🔗 Search Kerala High Court]({kerala_link})")

        if st.session_state.research_court in ["Supreme Court", "Both"]:
            sc_link = generate_google_link(
                "sci.gov.in",
                phrase,
                st.session_state.research_period
            )
            st.markdown(f"[🔗 Search Supreme Court]({sc_link})")

        st.markdown("---")

# ---------------------------------------------------
# VERIFIED JUDGMENT ENTRY
# ---------------------------------------------------
st.markdown("## 📥 Add Verified Judgment")

case_title = st.text_input("Case Title (Required)")
citation = st.text_input("Citation (Optional)")
extract = st.text_area("Relevant Extract from Official Judgment (Required)", height=150)

if st.button("➕ Add to Draft References"):
    if not case_title.strip() or not extract.strip():
        st.warning("Case Title and Extract are mandatory.")
    else:
        st.session_state.selected_references.append({
            "title": case_title.strip(),
            "citation": citation.strip(),
            "extract": extract.strip()
        })
        st.success("Judgment added.")

if st.session_state.selected_references:
    st.markdown("### 📚 Added References")

    for i, ref in enumerate(st.session_state.selected_references):
        with st.expander(ref["title"]):
            st.write(f"**Citation:** {ref['citation']}")
            st.write(ref["extract"][:500] + "...")
            if st.button(f"❌ Remove {i}", key=f"remove_{i}"):
                st.session_state.selected_references.pop(i)
                st.rerun()

# ---------------------------------------------------
# 9. ACTION BUTTONS
# ---------------------------------------------------
b1, b2, b3 = st.columns(3)

with b1:
    if st.button("🚀 Draft Standard", type="primary", use_container_width=True):

        references_text = ""
        if st.session_state.selected_references:
            references_text = "\nVerified Judgments:\n"
            for ref in st.session_state.selected_references:
                references_text += f"""
Case: {ref['title']}
Citation: {ref['citation']}
Extract:
{ref['extract']}
"""

        prompt = f"""
Draft {dtype} for {court} at {target_dist}.

Facts:
{st.session_state.facts_input}

IMPORTANT:
Use ONLY the verified judgment extracts below.
Do NOT create or assume any case law.
If no verified extract is provided, do not cite case law.

{references_text}

STRICT RULES:
- STRICTLY use PARTY A and PARTY B
"""

        with st.spinner("AI Drafting..."):
            res, tank, sec = smart_rotate_draft(prompt, st.session_state.facts_input, st.session_state.selected_model)

            if res:
                st.session_state.final_master = res
                st.session_state.draft_history.insert(
                    0,
                    {"label": f"{dtype} ({datetime.now().strftime('%H:%M')})", "content": res}
                )
                st.toast(f"Draft generated in {sec}s")

with b2:
    selected_ref = st.selectbox("Mirror Reference", ["None"] + os.listdir(VAULT_PATH))
    if st.button("✨ Mirror Style", use_container_width=True, disabled=(selected_ref == "None")):
        doc = Document(os.path.join(VAULT_PATH, selected_ref))
        dna = "\n".join([p.text for p in doc.paragraphs[:15]])
        prompt = f"Style DNA:\n{dna}\n\nDraft {dtype} for {court} at {target_dist}. Use PARTY A/B."
        with st.spinner("Mirroring..."):
            res, tank, sec = smart_rotate_draft(prompt, st.session_state.facts_input, st.session_state.selected_model)
            if res:
                st.session_state.final_master = res

with b3:
    if st.button("🗑️ Reset All", use_container_width=True):
        preserved_auth = st.session_state.get("authenticated", False)
        preserved_role = st.session_state.get("user_role", "user")
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.session_state.authenticated = preserved_auth
        st.session_state.user_role = preserved_role
        st.rerun()

# ---------------------------------------------------
# 10. STYLE VAULT
# ---------------------------------------------------
with st.expander("📁 Style Vault Upload"):
    uploaded = st.file_uploader("Upload Reference (.docx)", type="docx")
    if uploaded:
        with open(os.path.join(VAULT_PATH, uploaded.name), "wb") as f:
            f.write(uploaded.getbuffer())
        st.success("Uploaded successfully.")

# ---------------------------------------------------
# 11. DRAFT HISTORY
# ---------------------------------------------------
with st.expander("📜 Draft History (Last 10)"):
    for i, item in enumerate(st.session_state.draft_history[:10]):
        if st.button(item["label"], key=f"h_{i}"):
            st.session_state.final_master = item["content"]
            st.rerun()

# ---------------------------------------------------
# 12. EDITOR & DOWNLOAD
# ---------------------------------------------------
if st.session_state.final_master:

    st.divider()

    col_a, col_b, col_c = st.columns(3)

    with col_a:
        p_a = st.text_input("Petitioner Name:", key="pet_name")
        st.button("Map 'PARTY A'", on_click=perform_replacement, args=("PARTY A", p_a), use_container_width=True)

    with col_b:
        p_b = st.text_input("Respondent Name:", key="res_name")
        st.button("Map 'PARTY B'", on_click=perform_replacement, args=("PARTY B", p_b), use_container_width=True)

    with col_c:
        f_old = st.text_input("Find:", key="f_txt")
        f_new = st.text_input("Replace:", key="r_txt")
        st.button("Replace All", on_click=perform_replacement, args=(f_old, f_new), use_container_width=True)

    st.text_area("Live Editor", value=st.session_state.final_master, height=500, key="main_editor")

    d1, d2 = st.columns(2)

    with d1:
        doc_gen = Document()
        doc_gen.add_paragraph(st.session_state.final_master)
        bio = io.BytesIO()
        doc_gen.save(bio)
        st.download_button("📥 MS Word", data=bio.getvalue(), file_name=f"{dtype}.docx")

    with d2:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 10, st.session_state.final_master.encode('latin-1','replace').decode('latin-1'))
        st.download_button("📥 PDF", data=pdf.output(dest='S').encode('latin-1'), file_name=f"{dtype}.pdf")
